from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.4

PURPLE = RGBColor(0x6c, 0x34, 0x83)
DARK_PURPLE = RGBColor(0x5b, 0x2c, 0x6f)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF8, 0xF4, 0xFC)
RED = RGBColor(0xE7, 0x4C, 0x3C)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PURPLE if level == 1 else DARK_PURPLE
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        set_cell_shading(cell, '6C3483')
    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(row_cells[i], 'F8F4FC')
    doc.add_paragraph()
    return table

def add_info_box(text, prefix="INFO"):
    p = doc.add_paragraph()
    run = p.add_run(f"  {prefix}: ")
    run.bold = True
    run.font.color.rgb = PURPLE
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_warning_box(text):
    add_info_box(text, "WARNING")

def add_danger_box(text):
    p = doc.add_paragraph()
    run = p.add_run(f"  CRITICAL: ")
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_PURPLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)

# ═══════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SaatSaheli')
run.font.size = Pt(42)
run.font.color.rgb = PURPLE
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Technical Support Document')
run.font.size = Pt(22)
run.font.color.rgb = DARK_PURPLE

doc.add_paragraph()

for line in [
    'Version: 1.0',
    'Date: March 25, 2026',
    'Platform: Web Application (React + Spring Boot)',
    'Database: PostgreSQL (Neon Cloud)',
    'AI Engine: Hugging Face Inference API (Stable Diffusion XL)',
    'Prepared by: SaatSaheli Development Team',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.color.rgb = GREY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('INTERNAL & OPERATIONS USE')
run.font.size = Pt(11)
run.font.color.rgb = PURPLE
run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════
add_heading_styled('Table of Contents', 1)
toc_items = [
    '1. Platform Overview',
    '2. System Architecture',
    '3. Technical Support — Server & Deployment',
    '4. User Support — Account & Access Issues',
    '5. Database Support — PostgreSQL / Neon',
    '6. Admin Support — Admin Panel Operations',
    '7. Super Admin Support — Elevated Operations',
    '8. AI & Image Generation Support',
    '9. File Upload & Storage Support',
    '10. Chat & Community Support',
    '11. Payments & Subscription Support (Stripe)',
    '12. Security & Authentication',
    '13. API Reference — Quick Lookup',
    '14. Error Code Reference',
    '15. Environment Variables Reference',
    '16. Troubleshooting Guide',
    '17. Escalation Matrix & Contact',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_PURPLE

doc.add_page_break()

# ═══════════════════════════════════════
# 1. PLATFORM OVERVIEW
# ═══════════════════════════════════════
add_heading_styled('1. Platform Overview', 1)
doc.add_paragraph(
    'SaatSaheli (meaning "Seven Friends") is a full-stack creative publishing and community platform '
    'that enables users to create, publish, and share digital books, articles, and image galleries. '
    'The platform features AI-powered content generation with multilingual (Hindi/English) support.'
)

add_heading_styled('1.1 Key Features', 2)
features = [
    ('Book Publishing: ', 'Create, edit, and publish digital books with drag-and-drop page editing'),
    ('AI Cover Generation: ', 'Generate book covers using Stable Diffusion XL with 12+ artistic styles'),
    ('Multilingual AI Prompts: ', 'Hindi-to-English translation for AI image prompts (Helsinki-NLP)'),
    ('Document Upload: ', 'Import PDF and Word documents, auto-converted to page-based books'),
    ('Flipbook Reader: ', 'Interactive page-flip style book reading experience'),
    ('Galleries & Articles: ', 'Create and manage image galleries and articles'),
    ('Community Chat: ', 'Real-time chat rooms across 6 categories'),
    ('Social Features: ', 'Likes, comments, favorites, following/followers'),
    ('Subscription Plans: ', 'Free, Premium, Gold, and Creator tiers via Stripe'),
    ('Admin Dashboard: ', 'User management, content moderation, platform statistics'),
]
for bold, text in features:
    add_bullet(text, bold)

add_heading_styled('1.2 Technology Stack', 2)
add_table(
    ['Layer', 'Technology', 'Version'],
    [
        ['Frontend', 'React, React Router, Axios', '19.x'],
        ['Backend', 'Java, Spring Boot, Spring Data JPA', 'Java 17, Spring Boot 3.4.5'],
        ['Database', 'PostgreSQL (Neon Cloud)', 'Latest'],
        ['AI/ML', 'Hugging Face Inference API', 'Stable Diffusion XL, Helsinki-NLP'],
        ['File Storage', 'Cloudinary', 'v1.38.0'],
        ['Payments', 'Stripe', 'v28.2.0'],
        ['Auth', 'JWT (HS256), Google OAuth, BCrypt', 'jjwt 0.12.6'],
        ['Deployment', 'GitHub Pages (Frontend), Spring Boot JAR (Backend)', '—'],
    ]
)

# ═══════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════
add_heading_styled('2. System Architecture', 1)

add_heading_styled('2.1 High-Level Architecture', 2)
add_code_block(
    '  [React Frontend :3000]\n'
    '         |\n'
    '         | REST API (JSON over HTTPS)\n'
    '         |\n'
    '  [Spring Boot Backend :8081]\n'
    '         |\n'
    '    +--------+----------+---------+---------+\n'
    '    |        |          |         |         |\n'
    '  [Neon   [Cloudinary] [Stripe] [Hugging  [Google\n'
    '  Postgres]                      Face API]  OAuth]'
)

add_heading_styled('2.2 Backend Package Structure', 2)
add_code_block(
    'com.SaatSaheli.spring/\n'
    '  ├── config/         CorsConfig, JwtInterceptor, GlobalExceptionHandler\n'
    '  ├── controller/     AuthController, BookController, AdminController,\n'
    '  │                   ChatController, GalleryController, ArticleController,\n'
    '  │                   SocialController, PaymentController, HealthController\n'
    '  ├── model/          User, Login, Book, Page, Article, Gallery,\n'
    '  │                   GalleryImage, ChatRoom, ChatMessage, Comment,\n'
    '  │                   ContentLike, Favorite\n'
    '  ├── repository/     JPA repositories for all models\n'
    '  ├── service/        Business logic services\n'
    '  └── util/           RoleUtil, RateLimiter, JwtUtil'
)

add_heading_styled('2.3 Port Configuration', 2)
add_table(
    ['Service', 'Port', 'Notes'],
    [
        ['Frontend (React)', '3000', 'Development server via npm start'],
        ['Backend (Spring Boot)', '8081', 'Started via ./mvnw clean spring-boot:run'],
        ['PostgreSQL (Neon)', '5432', 'Cloud-hosted, SSL required'],
    ]
)

# ═══════════════════════════════════════
# 3. TECHNICAL SUPPORT
# ═══════════════════════════════════════
add_heading_styled('3. Technical Support — Server & Deployment', 1)

add_heading_styled('3.1 Starting the Backend', 2)
add_code_block('cd SaatSaheliMain/SaatSaheli/\n./mvnw clean spring-boot:run')
add_warning_box('Always use ./mvnw clean spring-boot:run when Java source files change. Incremental compilation can miss new files.')

add_heading_styled('3.2 Starting the Frontend', 2)
add_code_block('cd SaatSaheliMain/FrontEnd/\nnpm start')

add_heading_styled('3.3 Health Check', 2)
add_table(
    ['Endpoint', 'Method', 'Expected Response'],
    [['GET /api/health', 'GET', '{"status": "ok"}']]
)
doc.add_paragraph('Use this endpoint for uptime monitoring and load balancer health probes.')

add_heading_styled('3.4 Build & Deployment Notes', 2)
add_bullet('requires Node.js 20 LTS (build hangs on Node v25)', 'Frontend build: ')
add_bullet('GitHub Pages via gh-pages package', 'Frontend hosting: ')
add_bullet('./mvnw clean package produces executable JAR', 'Backend build: ')
add_bullet('Hibernate auto-migrates schema on startup (ddl-auto: update)', 'JPA DDL strategy: ')

add_heading_styled('3.5 Server Restart Checklist', 2)
steps = [
    'Verify .env file exists at SaatSaheli/.env with all required variables',
    'Run ./mvnw clean spring-boot:run',
    'Confirm health check returns {"status": "ok"}',
    'Verify database connectivity via any GET endpoint',
    'Test frontend connectivity to backend',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

# ═══════════════════════════════════════
# 4. USER SUPPORT
# ═══════════════════════════════════════
add_heading_styled('4. User Support — Account & Access Issues', 1)

add_heading_styled('4.1 User Registration', 2)
add_table(
    ['Endpoint', 'Method', 'Rate Limit'],
    [['/api/auth/signup', 'POST', '10 attempts per 15 minutes (per IP)']]
)
doc.add_paragraph('Required fields: name, email, password. Defaults: Role = USER, Plan = Free.')

add_heading_styled('4.2 User Login', 2)
add_table(
    ['Endpoint', 'Method', 'Auth Providers'],
    [['/api/auth/login', 'POST', 'Email/Password, Google OAuth']]
)
doc.add_paragraph('Returns a JWT token valid for 24 hours.')

add_heading_styled('4.3 Common User Issues & Resolutions', 2)
add_table(
    ['Issue', 'Cause', 'Resolution'],
    [
        ['Cannot sign up', 'Email already exists', 'Use a different email or reset password'],
        ['Cannot sign up', 'Rate limit exceeded (429)', 'Wait 15 minutes and retry'],
        ['Cannot log in', 'Wrong password', 'Use forgot password flow'],
        ['Cannot log in', 'Account blocked/disabled', 'Contact admin to reactivate'],
        ['Token expired', 'JWT expired after 24h', 'Log in again to get a new token'],
        ['Google login fails', 'OAuth misconfiguration', 'Verify Google OAuth client ID in frontend config'],
    ]
)

add_heading_styled('4.4 Password Reset Flow', 2)
doc.add_paragraph('User-Initiated Reset:', style='List Bullet')
doc.add_paragraph('1. User calls POST /api/auth/forgot-password with email')
doc.add_paragraph('2. System generates a temporary password (logged server-side)')
doc.add_paragraph('3. User calls POST /api/auth/reset-password with old + new password')
doc.add_paragraph()
doc.add_paragraph('Admin-Initiated Reset (Super Admin Only):', style='List Bullet')
doc.add_paragraph('1. Super Admin calls PUT /api/auth/admin-reset-password/{userId}')
doc.add_paragraph('2. Password is force-reset without requiring old password verification')

add_heading_styled('4.5 Account Statuses', 2)
add_table(
    ['Status', 'Description', 'User Can Login?'],
    [
        ['ACTIVE', 'Normal active account', 'Yes'],
        ['INACTIVE', 'Temporarily deactivated', 'No'],
        ['DISABLED', 'Administratively disabled', 'No'],
        ['BLOCKED', 'Blocked for policy violations', 'No'],
        ['DELETED', 'Soft-deleted account', 'No'],
    ]
)

# ═══════════════════════════════════════
# 5. DATABASE SUPPORT
# ═══════════════════════════════════════
add_heading_styled('5. Database Support — PostgreSQL / Neon', 1)

add_heading_styled('5.1 Connection Details', 2)
add_table(
    ['Parameter', 'Value'],
    [
        ['Provider', 'Neon Cloud PostgreSQL'],
        ['Connection URL', '${NEON_DB_URL} (env variable)'],
        ['Driver', 'org.postgresql.Driver'],
        ['ORM', 'Spring Data JPA / Hibernate'],
        ['DDL Strategy', 'update (auto-migrate)'],
    ]
)

add_heading_styled('5.2 Connection Pool Configuration (HikariCP)', 2)
add_table(
    ['Setting', 'Value', 'Reason'],
    [
        ['Maximum Pool Size', '5', 'Neon limits concurrent connections'],
        ['Connection Timeout', '20,000 ms (20s)', 'Allow time for cold-start connections'],
        ['Max Lifetime', '600,000 ms (10 min)', 'Neon closes long-lived idle connections'],
        ['Idle Timeout', '300,000 ms (5 min)', 'Release unused connections promptly'],
        ['Keep-Alive Time', '60,000 ms (1 min)', 'Prevent premature idle closure'],
    ]
)

add_heading_styled('5.3 Database Schema (Entity Tables)', 2)
add_table(
    ['Table', 'Description', 'Key Relationships'],
    [
        ['users', 'User profiles, roles, plans', 'One-to-one with logins'],
        ['logins', 'Credentials, auth provider, status', 'Linked to users'],
        ['books', 'Book metadata, status, category', 'Owned by user; has many pages'],
        ['pages', 'Book page content and numbering', 'Belongs to book'],
        ['articles', 'Articles with headline and content', 'Owned by user'],
        ['galleries', 'Image gallery containers', 'Owned by user; has many images'],
        ['gallery_images', 'Individual gallery images', 'Belongs to gallery'],
        ['chat_rooms', 'Chat room definitions', 'Has many messages'],
        ['chat_messages', 'Individual chat messages', 'Belongs to room; sent by user'],
        ['comments', 'Comments on content', 'Polymorphic by target type'],
        ['content_likes', 'Like toggles on content', 'Polymorphic by target type'],
        ['favorites', 'User-favorited content', 'Linked to user'],
    ]
)

add_heading_styled('5.4 Common Database Issues', 2)
add_table(
    ['Issue', 'Cause', 'Resolution'],
    [
        ['Connection refused', 'Neon instance sleeping', 'First request wakes it; retry after 5-10s'],
        ['Too many connections', 'Pool exhaustion (max 5)', 'Check for connection leaks; restart backend'],
        ['Connection reset', 'Neon closed idle connection', 'HikariCP auto-recovers; keepalive prevents this'],
        ['Null byte error (0x00)', 'Content contains null characters', 'Sanitize data before insert (strip \\u0000)'],
        ['Schema out of sync', 'Manual DB changes vs JPA model', 'Restart backend — Hibernate re-syncs'],
    ]
)

add_heading_styled('5.5 Database Maintenance Commands', 2)
add_code_block(
    '-- Check active connections\n'
    'SELECT count(*) FROM pg_stat_activity WHERE datname = \'your_db_name\';\n\n'
    '-- Check table sizes\n'
    'SELECT relname, pg_size_pretty(pg_total_relation_size(relid))\n'
    'FROM pg_catalog.pg_statio_user_tables ORDER BY pg_total_relation_size(relid) DESC;\n\n'
    '-- Check row counts\n'
    'SELECT schemaname, relname, n_live_tup\n'
    'FROM pg_stat_user_tables ORDER BY n_live_tup DESC;'
)

add_heading_styled('5.6 Data Migration History', 2)
add_info_box(
    'Migration completed March 2026: Migrated from Google Sheets to Neon PostgreSQL. '
    'Runner: DataMigrationRunner.java (run with -Dspring-boot.run.profiles=migrate). '
    'Migrated: 5 users, 5 logins, 22 books, 244 pages, 6 chat rooms, 4 chat messages.'
)

# ═══════════════════════════════════════
# 6. ADMIN SUPPORT
# ═══════════════════════════════════════
add_heading_styled('6. Admin Support — Admin Panel Operations', 1)

add_heading_styled('6.1 Admin Role Capabilities', 2)
doc.add_paragraph('Users with ADMIN role can perform the following operations:')
add_table(
    ['Operation', 'Endpoint', 'Method'],
    [
        ['View all users with activity stats', '/api/admin/users', 'GET'],
        ['Change user account status', '/api/admin/users/{userId}/status', 'PUT'],
        ['View all books', '/api/admin/books', 'GET'],
        ['Delete a book', '/api/admin/books/{bookId}', 'DELETE'],
        ['Archive a book', '/api/admin/books/{bookId}/archive', 'PUT'],
        ['Recover a deleted/archived book', '/api/admin/books/{bookId}/recover', 'PUT'],
        ['View platform statistics', '/api/admin/stats', 'GET'],
        ['Delete any chat message', '/api/chat/messages/{messageId}', 'DELETE'],
    ]
)

add_heading_styled('6.2 Platform Statistics Dashboard', 2)
doc.add_paragraph('The /api/admin/stats endpoint returns:')
add_bullet('Total users, active users, blocked users')
add_bullet('Total books, published books, draft books')
add_bullet('Creator vs. Visitor user counts')
add_bullet('Content activity metrics')

add_heading_styled('6.3 User Management Actions', 2)
doc.add_paragraph('Blocking a User:')
add_code_block('PUT /api/admin/users/{userId}/status\nHeader: Authorization: Bearer {admin_jwt_token}\nBody: { "status": "BLOCKED" }')
doc.add_paragraph('Reactivating a User:')
add_code_block('PUT /api/admin/users/{userId}/status\nHeader: Authorization: Bearer {admin_jwt_token}\nBody: { "status": "ACTIVE" }')

add_heading_styled('6.4 Content Moderation', 2)
add_bullet('Hides book from public view, can be recovered', 'Archive: ')
add_bullet('Soft-deletes the book (status = DELETED)', 'Delete: ')
add_bullet('Restores a deleted/archived book to DRAFT status', 'Recover: ')
add_bullet('Admin can delete any chat message (soft delete — shows as "[deleted]")', 'Chat moderation: ')

# ═══════════════════════════════════════
# 6.5 ADMIN vs SUPER ADMIN COMPARISON
# ═══════════════════════════════════════
add_heading_styled('6.5 Admin vs Super Admin — Key Differences', 2)
doc.add_paragraph(
    'The following table provides a clear side-by-side comparison of what Admin and Super Admin roles '
    'can and cannot do. This is useful for understanding the privilege boundary between the two roles.'
)
add_table(
    ['Capability', 'Admin', 'Super Admin'],
    [
        ['View all users with activity stats', 'Yes', 'Yes'],
        ['Change user account status (block/disable/activate)', 'Yes', 'Yes'],
        ['View / delete / archive / recover books', 'Yes', 'Yes'],
        ['View platform statistics', 'Yes', 'Yes'],
        ['Delete any chat message', 'Yes', 'Yes'],
        ['Change user roles', 'No', 'Yes'],
        ['Force password reset (no old password needed)', 'No', 'Yes'],
        ['Purge a single book permanently', 'No', 'Yes'],
        ['Purge ALL deleted books permanently', 'No', 'Yes'],
    ]
)

doc.add_paragraph('Summary of Key Differences:')
add_bullet(
    'Only Super Admin can promote/demote users (e.g., make someone an Admin). Admin cannot assign roles at all.',
    'Role Management — '
)
add_bullet(
    'Super Admin can reset any user\'s password without knowing the old one (/api/auth/admin-reset-password/{userId}). Admin cannot.',
    'Force Password Reset — '
)
add_bullet(
    'Admin can only soft-delete (status = DELETED) or archive books. Super Admin can permanently purge data from the database — this is irreversible.',
    'Permanent Data Deletion (Purge) — '
)
doc.add_paragraph()
doc.add_paragraph('In short: Admin handles routine moderation, while Super Admin handles sensitive, irreversible operations like role changes and permanent deletions.')

# ═══════════════════════════════════════
# 6.6 SUPPORT QUERIES — DISPLAY ID CONVENTION
# ═══════════════════════════════════════
add_heading_styled('6.6 Support Queries — Display ID Convention', 2)
doc.add_paragraph(
    'All public form submissions (Magazine Submission, Help & Support, Feedback, and Contact Us) '
    'are persisted to the contact_messages table and surfaced in the Admin Dashboard under the '
    '"Support Queries" tab. To make rows easier to reference verbally and in follow-up emails, the '
    'admin UI displays a prefixed, zero-padded ID derived from the form type and the database row id. '
    'The database primary key itself is unchanged — this is purely a display convention.'
)
add_table(
    ['Form', 'Subject Pattern', 'Display ID Prefix', 'Example'],
    [
        ['Magazine Submission', 'starts with "Magazine Submission:"', 'M',  'M00026'],
        ['Help & Support',      'starts with "Help & Support:"',      'HS', 'HS00027'],
        ['Feedback',            'starts with "feedback" (case-insensitive)', 'FE', 'FE00028'],
        ['Contact Us',          'anything else (fallback)',            'CU', 'CU00029'],
    ]
)
add_bullet(
    'A separate, randomly-generated tracking ID (e.g. SS-MAG-A4F2K9) is also issued at submission time '
    'and shown to the user on the success screen. Users can quote either the tracking ID or the display ID '
    'in any follow-up email.',
    'Tracking ID vs Display ID — '
)
add_bullet(
    'The Support Queries search box matches against name, email, subject, tracking ID, and the formatted '
    'display ID, so admins can paste any of these to locate a row.',
    'Search — '
)
add_bullet(
    'Classification mirrors EmailService.sendContactNotification, so the prefix shown in the table always '
    'matches the form type used to label admin notification emails.',
    'Consistency — '
)
doc.add_paragraph()

# ═══════════════════════════════════════
# 7. SUPER ADMIN SUPPORT
# ═══════════════════════════════════════
add_heading_styled('7. Super Admin Support — Elevated Operations', 1)
add_danger_box('SUPER_ADMIN operations are irreversible in some cases. Exercise extreme caution with purge operations.')

add_heading_styled('7.1 Super Admin Exclusive Capabilities', 2)
add_table(
    ['Operation', 'Endpoint', 'Method', 'Risk Level'],
    [
        ['Change user role', '/api/admin/users/{userId}/role', 'PUT', 'Medium'],
        ['Force password reset', '/api/auth/admin-reset-password/{userId}', 'PUT', 'Medium'],
        ['Purge single book permanently', '/api/admin/books/{bookId}/purge', 'DELETE', 'HIGH'],
        ['Purge all deleted books', '/api/admin/books/purge', 'DELETE', 'CRITICAL'],
    ]
)

add_heading_styled('7.2 Role Hierarchy', 2)
add_code_block(
    'SUPER_ADMIN  (highest - can do everything)\n'
    '    |\n'
    '  ADMIN      (can manage users and content, cannot assign roles)\n'
    '    |\n'
    '   USER      (standard user - create and manage own content)'
)

add_heading_styled('7.3 Role Assignment Rules', 2)
add_bullet('Only SUPER_ADMIN can assign or change user roles')
add_bullet('Valid roles: USER, ADMIN, SUPER_ADMIN')
add_bullet('ADMIN cannot promote other users to ADMIN or SUPER_ADMIN')
doc.add_paragraph('Assigning a Role:')
add_code_block('PUT /api/admin/users/{userId}/role\nHeader: Authorization: Bearer {superadmin_jwt_token}\nBody: { "role": "ADMIN" }')

add_heading_styled('7.4 Purge Operations', 2)
add_danger_box('Purge = Permanent Deletion. Data cannot be recovered after purge. Always verify before executing.')
add_bullet('Permanently deletes one book and all its pages', 'Single purge: ')
add_bullet('Permanently deletes ALL books with status DELETED', 'Bulk purge: ')

# ═══════════════════════════════════════
# 8. AI & IMAGE GENERATION
# ═══════════════════════════════════════
add_heading_styled('8. AI & Image Generation Support', 1)

add_heading_styled('8.1 AI Image Generation', 2)
add_table(
    ['Parameter', 'Details'],
    [
        ['Endpoint', 'POST /api/generate-image'],
        ['Model', 'Stable Diffusion XL Base 1.0'],
        ['Provider', 'Hugging Face Inference API'],
        ['API URL', 'https://router.huggingface.co/hf-inference/models/'],
        ['Auth', 'Bearer token via HUGGINGFACE_API_TOKEN'],
    ]
)

add_heading_styled('8.2 Available Style Presets', 2)
add_table(
    ['Style', 'Prompt Prefix'],
    [
        ['General', '"High quality, detailed illustration:"'],
        ['Children', '"Children\'s book illustration, colorful, whimsical style:"'],
        ['Poetry', '"Artistic, dreamy, poetic watercolor illustration:"'],
        ['Story', '"Vivid storytelling illustration, narrative scene:"'],
        ['Art', '"Fine art painting, museum quality, expressive brushwork:"'],
        ['Fantasy', '"Epic fantasy art, magical, otherworldly scene:"'],
        ['Realistic', '"Photorealistic, high detail, lifelike rendering:"'],
        ['Science', '"Scientific illustration, accurate, educational diagram style:"'],
        ['Technology', '"Futuristic, sleek technology concept art:"'],
        ['History', '"Historical illustration, period-accurate, vintage style:"'],
        ['Geography', '"Beautiful landscape, geographical illustration, natural scenery:"'],
        ['Politics', '"Editorial illustration, bold political commentary style:"'],
    ]
)

add_heading_styled('8.3 Multilingual Support', 2)
add_info_box('Hindi Prompt Support: The system detects Devanagari script in user prompts and auto-translates to English using the Helsinki-NLP/opus-mt-hi-en model before sending to Stable Diffusion.')

add_heading_styled('8.4 Retry Logic', 2)
add_bullet('Up to 3 retry attempts with 20-second delays')
add_bullet('Triggers on HTTP 503 (model loading / cold start)')
add_bullet('Generated images are uploaded to Cloudinary for persistence')

add_heading_styled('8.5 Common AI Issues', 2)
add_table(
    ['Issue', 'Cause', 'Resolution'],
    [
        ['503 Service Unavailable', 'Model cold start', 'Auto-retries (up to 3x with 20s delay)'],
        ['Image not generated', 'HF API token expired', 'Regenerate token at huggingface.co/settings/tokens'],
        ['Poor image quality', 'Vague prompt', 'Use more specific, descriptive prompts'],
        ['Hindi prompt not working', 'Translation API failure', 'Falls back to original text; check HF token'],
    ]
)

# ═══════════════════════════════════════
# 9. FILE UPLOAD & STORAGE
# ═══════════════════════════════════════
add_heading_styled('9. File Upload & Storage Support', 1)

add_heading_styled('9.1 Upload Configuration', 2)
add_table(
    ['Setting', 'Value'],
    [
        ['Max File Size', '10 MB per file'],
        ['Max Request Size', '50 MB per request'],
        ['Supported Formats', 'PDF, DOCX, DOC, JPEG, PNG, GIF, WebP, BMP, TIFF'],
        ['Storage Provider', 'Cloudinary'],
    ]
)

add_heading_styled('9.2 Document Processing', 2)
add_bullet('PDFBox extracts text per page + renders pages as images (150 DPI)', 'PDF Upload: ')
add_bullet('Apache POI extracts text with 500-character page pagination', 'Word Upload: ')
add_bullet('EXIF metadata (GPS, camera info) is stripped before upload', 'Privacy: ')

add_heading_styled('9.3 Common Upload Issues', 2)
add_table(
    ['Issue', 'Resolution'],
    [
        ['413 Payload Too Large', 'Reduce file size below 10MB'],
        ['Unsupported format', 'Convert to a supported format'],
        ['Cloudinary upload fails', 'Check Cloudinary credentials in .env'],
        ['PDF pages missing', 'Check PDF is not password-protected'],
    ]
)

# ═══════════════════════════════════════
# 10. CHAT & COMMUNITY
# ═══════════════════════════════════════
add_heading_styled('10. Chat & Community Support', 1)

add_heading_styled('10.1 Default Chat Rooms', 2)
add_table(
    ['Room', 'Category', 'Purpose'],
    [
        ['Art', 'Art', 'Discuss visual arts and artwork'],
        ['Music', 'Music', 'Music discussions and sharing'],
        ['Writing', 'Writing', 'Creative writing and feedback'],
        ['Tech', 'Tech', 'Technology and tools'],
        ['Creativity', 'Creativity', 'General creative discussions'],
        ['Community', 'Community', 'Announcements and general talk'],
    ]
)

add_heading_styled('10.2 Chat Features', 2)
add_bullet('Real-time polling every 5 seconds')
add_bullet('Pagination support via afterId and limit parameters')
add_bullet('Soft-delete for messages (shows as "[deleted]")')
add_bullet('Message deletion: allowed for sender or admin')
add_bullet('Auto-initialization: rooms are created on first access if none exist')

# ═══════════════════════════════════════
# 11. PAYMENTS
# ═══════════════════════════════════════
add_heading_styled('11. Payments & Subscription Support (Stripe)', 1)

add_heading_styled('11.1 Subscription Plans', 2)
add_table(
    ['Plan', 'Description'],
    [
        ['Free', 'Basic access, limited features'],
        ['Premium', 'Enhanced features and content access'],
        ['Gold', 'Full platform access with priority support'],
        ['Creator', 'All features + publishing and monetization tools'],
    ]
)

add_heading_styled('11.2 Payment Flow', 2)
doc.add_paragraph('1. User selects plan on frontend')
doc.add_paragraph('2. Frontend calls backend to create Stripe Checkout Session')
doc.add_paragraph('3. User is redirected to Stripe for payment')
doc.add_paragraph('4. Stripe sends webhook event on payment completion')
doc.add_paragraph('5. Backend webhook handler updates user\'s plan')

add_heading_styled('11.3 Stripe Configuration', 2)
add_table(
    ['Variable', 'Purpose'],
    [
        ['STRIPE_SECRET_KEY', 'Server-side Stripe API key'],
        ['STRIPE_WEBHOOK_SECRET', 'Webhook signature verification'],
    ]
)

# ═══════════════════════════════════════
# 12. SECURITY
# ═══════════════════════════════════════
add_heading_styled('12. Security & Authentication', 1)

add_heading_styled('12.1 JWT Authentication', 2)
add_table(
    ['Parameter', 'Value'],
    [
        ['Algorithm', 'HS256 (HMAC-SHA256)'],
        ['Token Expiry', '24 hours'],
        ['Header', 'Authorization: Bearer {token}'],
        ['Fallback Header', 'X-User-Id (legacy support)'],
        ['Secret', '${JWT_SECRET} environment variable'],
    ]
)

add_heading_styled('12.2 Rate Limiting', 2)
add_table(
    ['Endpoint', 'Key', 'Limit', 'Window'],
    [
        ['Signup', 'IP Address', '10 attempts', '15 minutes'],
        ['Login', 'Email / IP', '10 attempts', '15 minutes'],
        ['Forgot Password', 'IP Address', '10 attempts', '15 minutes'],
    ]
)
add_info_box('Rate limits are stored in-memory and reset on server restart.')

add_heading_styled('12.3 Password Security', 2)
add_bullet('Passwords hashed with BCrypt (one-way hashing)')
add_bullet('Plain-text passwords are never stored or logged')
add_bullet('Password comparison via BCrypt matches() function')

add_heading_styled('12.4 CORS Configuration', 2)
add_table(
    ['Setting', 'Value'],
    [
        ['Allowed Origins', 'Configurable via CORS_ALLOWED_ORIGINS'],
        ['Default Origins', 'http://localhost:3000, http://localhost:8081'],
        ['Allowed Methods', 'GET, POST, PUT, DELETE, OPTIONS'],
        ['Credentials', 'Allowed'],
    ]
)

# ═══════════════════════════════════════
# 13. API REFERENCE
# ═══════════════════════════════════════
add_heading_styled('13. API Reference — Quick Lookup', 1)
add_table(
    ['Category', 'Base Path', 'Auth Required'],
    [
        ['Authentication', '/api/auth/*', 'Varies'],
        ['Books', '/api/books/*', 'Yes (most)'],
        ['File Upload', '/api/upload', 'Yes'],
        ['File Download', '/api/download/*', 'Yes'],
        ['AI Image Generation', '/api/generate-image', 'Yes'],
        ['Chat', '/api/chat/*', 'Yes'],
        ['Social (Likes/Comments)', '/api/social/*', 'Yes'],
        ['Galleries', '/api/galleries/*', 'Yes (most)'],
        ['Articles', '/api/articles/*', 'Yes (most)'],
        ['Payments', '/api/payments/*', 'Yes'],
        ['Admin', '/api/admin/*', 'ADMIN / SUPER_ADMIN'],
        ['Health Check', '/api/health', 'No'],
    ]
)

# ═══════════════════════════════════════
# 14. ERROR CODES
# ═══════════════════════════════════════
add_heading_styled('14. Error Code Reference', 1)
add_table(
    ['HTTP Code', 'Meaning', 'Common Causes'],
    [
        ['200', 'Success', 'Request completed normally'],
        ['201', 'Created', 'Resource created successfully'],
        ['400', 'Bad Request', 'Invalid input, missing fields, validation failure'],
        ['401', 'Unauthorized', 'Missing or invalid JWT token'],
        ['403', 'Forbidden', 'Insufficient role/permissions'],
        ['404', 'Not Found', 'Resource does not exist'],
        ['409', 'Conflict', 'Duplicate resource (e.g., email already registered)'],
        ['413', 'Payload Too Large', 'File exceeds 10MB upload limit'],
        ['429', 'Too Many Requests', 'Rate limit exceeded (10 per 15 min)'],
        ['500', 'Internal Server Error', 'Unexpected server error'],
        ['503', 'Service Unavailable', 'AI model loading (Hugging Face cold start)'],
    ]
)
doc.add_paragraph('Standard error response format:')
add_code_block('{\n  "error": "Human-readable error message"\n}')

# ═══════════════════════════════════════
# 15. ENV VARIABLES
# ═══════════════════════════════════════
add_heading_styled('15. Environment Variables Reference', 1)
add_table(
    ['Variable', 'Purpose', 'Required'],
    [
        ['NEON_DB_URL', 'PostgreSQL connection URL', 'Yes'],
        ['NEON_DB_USERNAME', 'Database username', 'Yes'],
        ['NEON_DB_PASSWORD', 'Database password', 'Yes'],
        ['JWT_SECRET', 'JWT signing secret key', 'Yes'],
        ['HUGGINGFACE_API_TOKEN', 'HF Inference API token', 'Yes'],
        ['CLOUDINARY_CLOUD_NAME', 'Cloudinary account name', 'Yes'],
        ['CLOUDINARY_API_KEY', 'Cloudinary API key', 'Yes'],
        ['CLOUDINARY_API_SECRET', 'Cloudinary API secret', 'Yes'],
        ['STRIPE_SECRET_KEY', 'Stripe API secret key', 'Yes'],
        ['STRIPE_WEBHOOK_SECRET', 'Stripe webhook verification', 'Yes'],
        ['CORS_ALLOWED_ORIGINS', 'Allowed frontend origins', 'Yes'],
        ['GOOGLE_CREDENTIALS_FILE', 'Path to Google service account JSON', 'Optional'],
        ['GOOGLE_DRIVE_FOLDER_ID', 'Google Drive folder ID', 'Optional'],
        ['GOOGLE_SHEETS_SPREADSHEET_ID', 'Legacy Google Sheets ID', 'Optional'],
        ['REACT_APP_API_URL', 'Backend URL for frontend', 'Yes (frontend)'],
    ]
)
doc.add_paragraph('Environment file location: SaatSaheli/.env (gitignored, loaded via spring-dotenv)')

# ═══════════════════════════════════════
# 16. TROUBLESHOOTING
# ═══════════════════════════════════════
add_heading_styled('16. Troubleshooting Guide', 1)

add_heading_styled('16.1 Backend Won\'t Start', 2)
add_table(
    ['Symptom', 'Check', 'Fix'],
    [
        ['Port 8081 in use', 'lsof -i :8081', 'Kill existing process or change port'],
        ['DB connection fails', 'Check .env variables', 'Verify Neon DB URL, username, password'],
        ['Missing .env file', 'Check SaatSaheli/.env exists', 'Create from template with required vars'],
        ['Java version mismatch', 'java -version', 'Install Java 17+'],
        ['Compilation error', 'Check build output', 'Use ./mvnw clean before build'],
    ]
)

add_heading_styled('16.2 Frontend Won\'t Start', 2)
add_table(
    ['Symptom', 'Check', 'Fix'],
    [
        ['Port 3000 in use', 'lsof -i :3000', 'Kill existing process'],
        ['Missing dependencies', 'Check node_modules/', 'Run npm install'],
        ['Build hangs', 'Check Node version', 'Use Node.js 20 LTS (not v25)'],
        ['API calls fail', 'Check REACT_APP_API_URL', 'Set correct backend URL'],
    ]
)

add_heading_styled('16.3 AI Image Generation Fails', 2)
add_table(
    ['Symptom', 'Check', 'Fix'],
    [
        ['503 errors persist', 'HF model status', 'Wait for model to load; check HF status page'],
        ['401 Unauthorized', 'HF API token', 'Regenerate token at huggingface.co'],
        ['Empty response', 'Cloudinary config', 'Verify Cloudinary credentials'],
    ]
)

add_heading_styled('16.4 Payment Issues', 2)
add_table(
    ['Symptom', 'Check', 'Fix'],
    [
        ['Checkout fails', 'Stripe key validity', 'Verify STRIPE_SECRET_KEY'],
        ['Webhook not received', 'Webhook URL config', 'Verify STRIPE_WEBHOOK_SECRET and endpoint URL'],
        ['Plan not updated', 'Server logs', 'Check webhook handler for errors'],
    ]
)

# ═══════════════════════════════════════
# 17. ESCALATION MATRIX
# ═══════════════════════════════════════
add_heading_styled('17. Escalation Matrix & Contact', 1)

add_heading_styled('17.1 Support Tiers', 2)
add_table(
    ['Tier', 'Scope', 'Handled By', 'Response Time'],
    [
        ['L1 — Basic', 'Login issues, password resets, general queries', 'Support Staff / Admin', 'Within 24 hours'],
        ['L2 — Technical', 'API errors, upload failures, AI issues, payment problems', 'Development Team', 'Within 48 hours'],
        ['L3 — Critical', 'Database issues, security incidents, server downtime', 'Senior Developer / Super Admin', 'Within 4 hours'],
        ['L4 — Emergency', 'Data breach, complete outage, payment data exposure', 'Project Lead + All Hands', 'Immediate'],
    ]
)

add_heading_styled('17.2 Escalation Triggers', 2)
add_bullet('Escalate to L2 if L1 cannot resolve within 24 hours or issue is technical')
add_bullet('Escalate to L3 if issue affects multiple users, involves data loss, or server is unresponsive')
add_bullet('Escalate to L4 if there is a confirmed security breach or complete platform outage')

add_heading_styled('17.3 Key Operational Contacts', 2)
add_table(
    ['Role', 'Responsibility'],
    [
        ['Super Admin', 'Role management, force resets, permanent data operations'],
        ['Admin', 'User management, content moderation, platform monitoring'],
        ['Developer', 'Bug fixes, API issues, deployment, infrastructure'],
    ]
)

add_heading_styled('17.4 Useful Monitoring Commands', 2)
add_code_block(
    '# Health check\n'
    'curl http://localhost:8081/api/health\n\n'
    '# Check backend process\n'
    'lsof -i :8081\n\n'
    '# Check frontend process\n'
    'lsof -i :3000\n\n'
    '# Database connection test\n'
    'psql "${NEON_DB_URL}" -c "SELECT 1;"'
)

# ── FOOTER ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SaatSaheli Technical Support Document | Version 1.0 | March 25, 2026')
run.font.size = Pt(9)
run.font.color.rgb = GREY
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('This document is for internal and operational use. Keep credentials and access details confidential.')
run.font.size = Pt(9)
run.font.color.rgb = GREY

# ── SAVE ──
output_path = '/Users/chitrasharma/Documents/GitHub/SaatSaheliMain/SaatSaheli_Support_Document.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
